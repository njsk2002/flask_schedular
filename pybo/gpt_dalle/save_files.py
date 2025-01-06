from fpdf import FPDF
from docx import Document
import pandas as pd
import matplotlib.pyplot as plt

class SaveFiles:
    @staticmethod
    def save_as_pdf(story, decisionQuestion, choices):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        pdf.cell(200, 10, txt="NovelGPT Output", ln=True, align='C')
        pdf.ln(10)

        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, txt=f"Story:\n{story}\n")
        pdf.ln(5)
        pdf.multi_cell(0, 10, txt=f"Decision Question:\n{decisionQuestion}\n")
        pdf.ln(5)

        pdf.cell(0, 10, txt="Choices:", ln=True)
        for choice in choices:
            pdf.cell(0, 10, txt=f"- {choice}", ln=True)

        pdf.output("NovelGPT_Output.pdf")


    @staticmethod
    def save_as_word(story, decisionQuestion, choices):
        doc = Document()
        doc.add_heading('NovelGPT Output', level=1)

        doc.add_heading('Story', level=2)
        doc.add_paragraph(story)

        doc.add_heading('Decision Question', level=2)
        doc.add_paragraph(decisionQuestion)

        doc.add_heading('Choices', level=2)
        for choice in choices:
            doc.add_paragraph(choice, style='List Bullet')

        doc.save("NovelGPT_Output.docx")
    
    @staticmethod
    def save_as_html(story, decisionQuestion, choices):
        with open("NovelGPT_Output.html", "w", encoding="utf-8") as file:
            file.write("<html><head><title>NovelGPT Output</title></head><body>")
            file.write("<h1>NovelGPT Output</h1>")
            file.write(f"<h2>Story</h2><p>{story}</p>")
            file.write(f"<h2>Decision Question</h2><p>{decisionQuestion}</p>")
            file.write("<h2>Choices</h2><ul>")
            for choice in choices:
                file.write(f"<li>{choice}</li>")
            file.write("</ul></body></html>")


    @staticmethod
    def save_as_excel(story, decisionQuestion, choices):
        # 데이터를 DataFrame으로 정리
        data = {
            "Category": ["Story", "Decision Question", "Choices"],
            "Content": [story, decisionQuestion, "\n".join(choices)]
        }
        df = pd.DataFrame(data)

        # Excel 파일로 저장
        file_path = "NovelGPT_Output.xlsx"
        with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name="NovelGPT")
            worksheet = writer.sheets["NovelGPT"]
            worksheet.set_column("A:A", 20)  # 카테고리 열 너비 설정
            worksheet.set_column("B:B", 50)  # 내용 열 너비 설정
        return file_path
    

    @staticmethod
    def save_as_image(story, decisionQuestion, choices, file_path="NovelGPT_Output.png"):
        # 이미지 생성
        fig, ax = plt.subplots(figsize=(8, 10))
        ax.axis("off")  # 축 제거

        # 텍스트 조합
        text = f"NovelGPT Output\n\nStory:\n{story}\n\nDecision Question:\n{decisionQuestion}\n\nChoices:\n"
        text += "\n".join([f"- {choice}" for choice in choices])

        # 텍스트 추가
        ax.text(0, 1, text, fontsize=10, va="top", wrap=True)
        plt.savefig(file_path, bbox_inches="tight")
        return file_path

